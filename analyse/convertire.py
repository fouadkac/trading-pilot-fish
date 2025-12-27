import json
from datetime import datetime
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox

# ----------- Fonctions de traitement -----------

def lire_tableau_docx(filepath):
    doc = Document(filepath)
    all_data = []

    if not doc.tables:
        raise ValueError("Le document ne contient aucun tableau.")

    table = doc.tables[0]  # on prend le premier tableau

    for row in table.rows:
        ligne = []
        for cell in row.cells:
            texte = cell.text.strip()
            ligne.append(texte)
        all_data.append([ligne])  # ajouter dimension supplémentaire

    return all_data

def supprimer_secondes(data):
    # Ignore la première ligne (en-têtes)
    for i in range(1, len(data)):
        line = data[i][0]
        try:
            original_time = line[0]
            dt = datetime.strptime(original_time, "%Y.%m.%d %H:%M:%S")
            line[0] = dt.strftime("%Y.%m.%d %H:%M")
        except Exception as e:
            print(f"Erreur à la ligne {i} : {e}")
    return data

def sauvegarder_json(data, filepath):
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4, ensure_ascii=False)

# ----------- Interface Tkinter -----------

class Application(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("Conversion DOCX -> JSON Professionnelle")
        self.geometry("600x250")
        self.configure(bg="#2E3440")  # fond sombre professionnel

        # Variables
        self.docx_path = tk.StringVar()
        self.json_path = tk.StringVar()

        # Widgets
        self.creer_widgets()

    def creer_widgets(self):
        lbl_title = tk.Label(self, text="Convertisseur DOCX (tableau) vers JSON 3D", 
                             font=("Segoe UI", 14, "bold"), bg="#2E3440", fg="#D8DEE9")
        lbl_title.pack(pady=(15, 10))

        # Fichier DOCX
        frame_docx = tk.Frame(self, bg="#2E3440")
        frame_docx.pack(pady=5, padx=20, fill='x')
        lbl_docx = tk.Label(frame_docx, text="Fichier DOCX :", bg="#2E3440", fg="#D8DEE9")
        lbl_docx.pack(side="left")
        ent_docx = tk.Entry(frame_docx, textvariable=self.docx_path, width=45)
        ent_docx.pack(side="left", padx=5)
        btn_browse_docx = tk.Button(frame_docx, text="Parcourir", command=self.ouvrir_docx, bg="#5E81AC", fg="white",
                                    activebackground="#81A1C1", relief="flat")
        btn_browse_docx.pack(side="left")

        # Fichier JSON destination
        frame_json = tk.Frame(self, bg="#2E3440")
        frame_json.pack(pady=5, padx=20, fill='x')
        lbl_json = tk.Label(frame_json, text="Fichier JSON :", bg="#2E3440", fg="#D8DEE9")
        lbl_json.pack(side="left")
        ent_json = tk.Entry(frame_json, textvariable=self.json_path, width=45)
        ent_json.pack(side="left", padx=5)
        btn_browse_json = tk.Button(frame_json, text="Enregistrer sous", command=self.enregistrer_json, bg="#5E81AC", fg="white",
                                   activebackground="#81A1C1", relief="flat")
        btn_browse_json.pack(side="left")

        # Bouton convertir
        btn_convert = tk.Button(self, text="Convertir et Nettoyer Heure", command=self.convertir,
                                bg="#A3BE8C", fg="#2E3440", font=("Segoe UI", 12, "bold"),
                                activebackground="#B5D99C", relief="flat")
        btn_convert.pack(pady=20)

    def ouvrir_docx(self):
        filename = filedialog.askopenfilename(filetypes=[("Fichiers Word", "*.docx")])
        if filename:
            self.docx_path.set(filename)
            # Propose un nom JSON par défaut
            if not self.json_path.get():
                default_json = filename.rsplit('.', 1)[0] + "_result.json"
                self.json_path.set(default_json)

    def enregistrer_json(self):
        filename = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("Fichiers JSON", "*.json")])
        if filename:
            self.json_path.set(filename)

    def convertir(self):
        docx_file = self.docx_path.get()
        json_file = self.json_path.get()

        if not docx_file:
            messagebox.showerror("Erreur", "Veuillez sélectionner un fichier DOCX.")
            return
        if not json_file:
            messagebox.showerror("Erreur", "Veuillez sélectionner un fichier JSON de destination.")
            return

        try:
            data_3d = lire_tableau_docx(docx_file)
            data_sans_secondes = supprimer_secondes(data_3d)
            sauvegarder_json(data_sans_secondes, json_file)
            messagebox.showinfo("Succès", f"Conversion terminée.\nFichier JSON sauvegardé dans:\n{json_file}")
        except Exception as e:
            messagebox.showerror("Erreur", f"Une erreur est survenue :\n{e}")

# ----------- Lancement -----------

if __name__ == "__main__":
    app = Application()
    app.mainloop()
